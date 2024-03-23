from docx import Document
import csv


students = []
with open("data/students.csv", "r", encoding="utf8") as f:
    for row in csv.DictReader(f):
        students.append(row)

accounts = []
with open("output/accounts.csv", "r", encoding="utf8") as f:
    for row in csv.DictReader(f):
        accounts.append(row)


def make_columns_bold(*columns):
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


document = Document()

for i, student in enumerate(students):
    sbd = student["sbd"]
    hoten = student["hoten"]
    ngaysinh = student["ngaysinh"]
    gioitinh = student["gioitinh"]
    donvi = student["donvi"]
    taikhoan = ""
    matkhau = ""

    for account in accounts:
        if account["username"] == ("u" + sbd):
            taikhoan = account["username"]
            matkhau = account["password"]

    table = document.add_table(rows=7, cols=2)

    header = table.columns[0]
    header.cells[0].text = "SBD:"
    header.cells[1].text = "Họ tên:"
    header.cells[2].text = "Ngày sinh:"
    header.cells[3].text = "Giới tính:"
    header.cells[4].text = "Đơn vị:"
    header.cells[5].text = "Tài khoản:"
    header.cells[6].text = "Mật khẩu:"
    make_columns_bold(header)

    column = table.columns[1]
    column.cells[0].text = sbd
    column.cells[1].text = hoten
    column.cells[2].text = ngaysinh
    column.cells[3].text = gioitinh
    column.cells[4].text = donvi
    column.cells[5].text = taikhoan
    column.cells[6].text = matkhau

    document.add_paragraph("\n")


# document.add_page_break()

document.save("output/demo.docx")
